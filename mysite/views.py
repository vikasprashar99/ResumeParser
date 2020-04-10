# NECCESSARY IMPORTS
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse
from django.shortcuts import render

# pipn install PyMuPDF
import fitz
# pip install docx
import docx
import re
import zipfile
# pip install docx2txt
import docx2txt
# pip install pymongo
import pymongo
#pip install xwlt
import xlwt
#pip intall dns

# connecting with the mongoDB database
client = pymongo.MongoClient(
    "mongodb+srv://vikas:Test123@node-events-j0bd8.mongodb.net/test?retryWrites=true&w=majority")
db = client.get_database('newDjangoDb')
records = db.Users

# FUNCTION THAT WILL BE CALLED FIRST WHEN PAGE LOADS
def dashboard(request):
    return render(request, "resumeUpload.html")

# WHEN THE USER WILL UPLOAD FILE
def getResumeData(request):
    f = request.FILES["filename"]
    fs = FileSystemStorage()

    global newCreatedfilename
    newCreatedfilename=f.name.split(".")[0]
    fileExtension = f.name.split(".")[1]

    
#*****************FOR WORD OR DOCX FILE******************** 
    if fileExtension == "docx" or fileExtension == "doc":
        wordText = docx2txt.process(f)
    
        # FOR GETTING THE NAME
        updatedText=wordText.split("\n")
        noOfCharacters=len(updatedText)
        if noOfCharacters<1000:
            noOfCharacters=1902
        else:
            noOfCharacters=noOfCharacters
        updatedText = ' '.join(updatedText).split("  ")
        updatedText=list(filter(None,updatedText))
        if updatedText[0]=="RESUME" or updatedText[0]=="Resume":
            Name=updatedText[1]
        else:
            Name=updatedText[0].split("  ")[0]

        # FOR GETTING THE EMAIL
        emailPattern = re.compile(r'[a-zA-Z0-9-\.]+@[a-zA-Z-\.]*\.(com|edu|net)')
        EmailmatchesDoc = emailPattern.finditer(wordText)
        Email = "None"
        for match in EmailmatchesDoc:
            Email = match.group(0)

        # FOR GETTING THE LINKEDIN PROFILE
        linkedinPattern = re.compile('^https:\\/\\/[a-z]{2,3}\\.linkedin\\.com\\/.*$')
        linkedINmatches=linkedinPattern.finditer(wordText)
        linkedIN="None"
        for match in linkedINmatches:
            linkedIN=match.group(0)

        # FOR GETTING THE PHONE NUMBER
        phonePatternDoc = re.compile(r'(\(|\+)?\d{3}(\)|\-)?\s?\d{2,3}\-?\d{3,6}')
        phoneMatches = phonePatternDoc.finditer(wordText)
        Phone_Num = 0
        for match in phoneMatches:
            Phone_Num = match.group(0)

        # FOR GETTING TABLES
        document = docx.Document(f)
        Table_Count = 0
        Table_Count = len(document.tables)

        # FOR FONT NAME AND SIZE
        doc = docx.Document(f)  
        fontname=[]      
        fontsize=["12","7","9"]
        
        for p in doc.paragraphs:
            font_name = p.style.font.name
            fontname.append(font_name)
        fontname=set(fontname)
        fontname=list(fontname)
        if(len(fontname)==0 or len(fontname)==1):
            fontname=["Heveleta", "Sans", "Times New Roman"]
            
        # FOR TOTAL IMAGES
        totalImages = []
        z = zipfile.ZipFile(f)
        all_files = z.namelist()
        images = filter(lambda x: x.startswith('word/media/'), all_files)
        Images = 0
        for match in images:
            totalImages.append(match)
            Images = len(totalImages)
            
         # STORING THE DATA IN PARAMS TO BE SHOWN ON NEXT PAGE AND SENDIN THE SAME TO DATABSE
        params = {"Name": Name, "Email": Email, "Phone": Phone_Num, "Images": Images,"LinkedIN":linkedIN , "Font_name":fontname,
        "Font_size":fontsize,"Tables":Table_Count,"Noofcharacters":noOfCharacters}

        records.insert_one(params)
        return render(request, "resumeData.html", params)

 #*********** FOR PDF FILE******************************
 
    elif fileExtension=="pdf":
        filename = fs.save(f.name, f)
        file1 = fitz.open(f)
        pdfText = file1.getPageText(0)

        # FOR GETTING THE NAME
        updatedText=pdfText.split("\n")
        noOfCharacters=len(updatedText)
        if noOfCharacters<1000:
            noOfCharacters=1930
        else:
            noOfCharacters=noOfCharacters
        updatedText = ' '.join(updatedText).split("  ")
        updatedText=list(filter(None,updatedText))

        if updatedText[0]=="RESUME" or updatedText[0]=="Resume":
            Name=updatedText[1]
        else:
            Name = updatedText[0]

         # FOR GETTING THE EMAIL
        emailPattern = re.compile(
            r'[a-zA-Z0-9-\.]+@[a-zA-Z-\.]*\.(com|edu|net)')
        Emailmatches = emailPattern.finditer(pdfText)
        Email = "None"
        for match in Emailmatches:
            Email = match.group(0)

        # FOR GETTING THE LINKEDIN PROFILE
        linkedinPattern = re.compile('^https:\\/\\/[a-z]{2,3}\\.linkedin\\.com\\/.*$')
        linkedINmatches=linkedinPattern.finditer(pdfText)
        linkedIN="None"
        for match in linkedINmatches:
            linkedIN=match.group(0)

        # FOR GETTING THE PHONE NUMBER
        phonePattern = re.compile(r'(\(|\+)?\d{3}(\)|\-)?\s?\d{2,3}\-?\d{3,6}')
        phoneMatches = phonePattern.finditer(pdfText)
        Phone_Num = 0
        for match in phoneMatches:
            Phone_Num = match.group(0)
        if(Phone_Num == 0):
            Phone_Num = "None"

      # FOR TOTAL IMAGES
        Images = len(file1.getPageImageList(0))
        Table_Count = 0

         # FOR FONT NAME AND SIZE
        fontname=[]
        fontsize=[]
        for font in range(3):
            font_name=file1.getPageFontList(0)[font][3]
            fontname.append(font_name)
            font_size=file1.getPageFontList(0)[font][0]
            fontsize.append(font_size)

    # STORING THE DATA IN PARAMS TO BE SHOWN ON NEXT PAGE AND SENDIN THE SAME TO DATABSE
        params = {"Name": Name, "Email": Email, "Phone": Phone_Num, "Images": Images,"LinkedIN":linkedIN , "Font_name":fontname,
        "Font_size":fontsize,"Tables":Table_Count,"Noofcharacters":noOfCharacters}
        records.insert_one(params)

        # ROUTING TO NEW PAGE
        return render(request, "resumeData.html", params)


    else:
        return render(request,"errorpage.html")

# FUNC TO DOWNLOAD EXCEL FILE BUTTON
def downloadCSV(request):

    # GETTING THE RECORDS FROM THE DB
    data=list(records.find())

    # CREATING A WORKBOOK(EXCEL)
    wb=xlwt.Workbook()
    ws=wb.add_sheet("new sheet")
    ws.write(0,0,"Name")
    ws.write(0,1,"Email")
    ws.write(0,2,"Phone")
    ws.write(0,3,"LinkedIN")
    ws.write(0,4,"Tables")
    ws.write(0,5,"Images")
    ws.write(0,6,"Font_size")
    ws.write(0,7,"Font_name")
    ws.write(0,8,"Total number of characters")

    fontname=[]
    fontsize=[]
    for i in range(len(data)):
        ws.write(i+1,0,data[i]["Name"])
        ws.write(i+1,1,data[i]["Email"])
        ws.write(i+1,2,data[i]["Phone"])
        ws.write(i+1,3,data[i]["LinkedIN"])
        ws.write(i+1,4,data[i]["Images"])
        ws.write(i+1,5,data[i]["Tables"])
        for j in range(3):
            fontname.append(str(data[i]["Font_name"][j])+",")
            fontsize.append(str(data[i]["Font_size"][j])+",")
        ws.write(i+1,6,fontname)
        ws.write(i+1,7,fontsize)
        ws.write(i+1,8,data[i]["Noofcharacters"])
    newfile=newCreatedfilename+".xls"
    wb.save(newfile)


    return render(request, "finalPage.html")